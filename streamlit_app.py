# app.py
# Streamlit Meeting Transcriber & Minutes Generator
# Features:
# - Live mic capture & rolling transcription (via streamlit-webrtc)
# - Record-to-file + transcribe
# - Upload audio + transcribe
# - Minutes generator (LLM) + download as .txt or .docx
#
# Backends:
#   1) OpenAI (API)  -> good accuracy, simple setup, chunked near-real-time
#   2) Local faster-whisper -> offline (CPU/GPU), best for uploads or post-call
#
# ------- HOW TO RUN -------
# 1) pip install -r requirements.txt  (see list after this code block)
# 2) streamlit run app.py
# 3) In the sidebar, choose your backend and (if OpenAI) paste your API key.

import os
import io
import time
import queue
import tempfile
from datetime import datetime

import streamlit as st
import streamlit_webrtc

# Audio / WebRTC
from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration, AudioProcessorBase
import av
import numpy as np
from pydub import AudioSegment

# Optional local transcription
try:
    from faster_whisper import WhisperModel
    HAVE_FASTER = True
except Exception:
    HAVE_FASTER = False

# Optional document export
try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# -------------------- UI SETUP --------------------
st.set_page_config(page_title="AI Meeting Minutes Generator", page_icon="📝", layout="wide")
st.title("📝 AI Meeting Minutes Generator")

with st.sidebar:
    st.header("⚙️ Settings")

    backend = st.selectbox(
        "Transcription backend",
        ["OpenAI (API)", "Local faster-whisper (offline)"],
        index=0
    )

    if backend == "OpenAI (API)":
        openai_key = st.text_input("OpenAI API key", type="password", placeholder="sk-...", help="Only stored in memory for this session.")
        # Lazy import to avoid dependency if unused
        if openai_key:
            os.environ["OPENAI_API_KEY"] = openai_key
        try:
            from openai import OpenAI
            oa_client = OpenAI(api_key=openai_key) if openai_key else None
        except Exception:
            oa_client = None
    else:
        model_size = st.selectbox("faster-whisper model", ["base", "small", "medium", "large-v3"], index=0)
        compute_type = st.selectbox("Compute type", ["int8", "int8_float16", "float16", "float32"], index=0)
        device = st.selectbox("Device", ["auto", "cpu", "cuda"], index=0)
        if HAVE_FASTER and st.button("Load / Warm up model"):
            st.session_state["_fw_model"] = WhisperModel(model_size, device=device if device!="auto" else None, compute_type=compute_type)
            st.success("faster-whisper model loaded.")

    st.divider()
    st.subheader("Minutes format")
    minutes_sections = st.multiselect(
        "Include sections",
        ["Attendees", "Agenda", "Key Points", "Decisions", "Action Items", "Risks/Issues", "Next Steps", "Summary"],
        default=["Attendees", "Key Points", "Decisions", "Action Items", "Next Steps", "Summary"]
    )

    st.caption("Tip: live mode sends short rolling chunks; for the best transcript, also click **Save full recording** and transcribe the final file.")

# Shared session state
if "rolling_transcript" not in st.session_state:
    st.session_state.rolling_transcript = []
if "full_audio" not in st.session_state:
    st.session_state.full_audio = AudioSegment.silent(duration=0)
if "segments" not in st.session_state:
    st.session_state.segments = []  # [(start, end, text)]
if "minutes_text" not in st.session_state:
    st.session_state.minutes_text = ""

# -------------------- TRANSCRIBERS --------------------
def transcribe_bytes_openai(wav_bytes: bytes, language: str | None = None) -> str:
    """
    Calls OpenAI's transcription on a small chunk. Returns plain text.
    """
    if not oa_client:
        raise RuntimeError("OpenAI client not initialized. Add API key in the sidebar.")
    # Use the "whisper-1" transcription model (or gpt-4o-mini-transcribe if enabled in your org)
    # This API expects a file-like object.
    resp = oa_client.audio.transcriptions.create(
        model="whisper-1",
        file=("chunk.wav", wav_bytes, "audio/wav"),
        language=language
    )
    return resp.text.strip()

def transcribe_file_openai(path: str, language: str | None = None) -> str:
    with open(path, "rb") as f:
        return transcribe_bytes_openai(f.read(), language)

def transcribe_file_faster(path: str, language: str | None = None) -> str:
    if "_fw_model" not in st.session_state:
        # Lazy-load if user didn't press the warm-up button
        if not HAVE_FASTER:
            raise RuntimeError("faster-whisper is not installed.")
        st.session_state["_fw_model"] = WhisperModel("base", device=None, compute_type="int8")
    model: WhisperModel = st.session_state["_fw_model"]
    segments, _ = model.transcribe(path, language=language, vad_filter=True)
    out = []
    st.session_state.segments = []  # reset with timestamps from this pass
    for seg in segments:
        out.append(seg.text)
        st.session_state.segments.append((seg.start, seg.end, seg.text))
    return " ".join(out).strip()

# -------------------- UTILITIES --------------------
def float_to_pcm16(audio: np.ndarray) -> np.ndarray:
    # audio expected in range [-1, 1]
    audio = np.clip(audio, -1.0, 1.0)
    return (audio * 32767.0).astype(np.int16)

def np_to_wav_bytes(audio_np: np.ndarray, sample_rate: int) -> bytes:
    # audio_np: (n_samples,) float32/float64 in [-1, 1]
    pcm16 = float_to_pcm16(audio_np)
    seg = AudioSegment(
        data=pcm16.tobytes(),
        sample_width=2,
        frame_rate=sample_rate,
        channels=1
    )
    buf = io.BytesIO()
    seg.export(buf, format="wav")
    return buf.getvalue()

def save_segment_to_session(seg_bytes: bytes):
    """Append chunk to the full session audio for final transcription/export."""
    new = AudioSegment.from_file(io.BytesIO(seg_bytes), format="wav")
    st.session_state.full_audio += new

def export_full_audio_wav() -> bytes:
    buf = io.BytesIO()
    st.session_state.full_audio.export(buf, format="wav")
    return buf.getvalue()

# -------------------- MINUTES GENERATOR --------------------
MINUTES_PROMPT = """You are an expert meeting minutes taker.
Given the raw transcript below, produce concise, structured minutes.
Use bullet points. Be specific with owners and dates when mentioned.
Sections to include (only output ones requested, in this order):
{sections}

Transcript:
\"\"\" 
{transcript}
\"\"\""""

def generate_minutes_llm(transcript: str, sections: list[str]) -> str:
    # Prefer OpenAI for summarization (quality + speed). You can switch to a local LLM if needed.
    if not oa_client:
        raise RuntimeError("OpenAI client not initialized. Add API key in the sidebar.")
    prompt = MINUTES_PROMPT.format(
        sections="\n".join(f"- {s}" for s in sections),
        transcript=transcript
    )
    chat = oa_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return chat.choices[0].message.content.strip()

# -------------------- LIVE MODE (WebRTC) --------------------
class AudioChunker(AudioProcessorBase):
    """
    Collects audio frames from WebRTC and yields rolling transcriptions
    by sending ~5s chunks. Also builds a full-session recording in memory.
    """
    SAMPLE_RATE = 16000
    CHANNELS = 1
    CHUNK_SECONDS = 5

    def __init__(self) -> None:
        self.buffer = np.zeros(0, dtype=np.float32)
        self.last_chunk_time = time.time()

    def recv_audio(self, frame: av.AudioFrame) -> av.AudioFrame:
        # Convert to mono float32 [-1, 1]
        buf = frame.to_ndarray().astype(np.float32) / 32768.0  # int16 -> float
        if buf.ndim > 1:
            buf = buf.mean(axis=0)  # mixdown to mono
        self.buffer = np.concatenate([self.buffer, buf])

        # Every CHUNK_SECONDS, cut a chunk and transcribe
        samples_per_chunk = int(self.SAMPLE_RATE * self.CHUNK_SECONDS)
        while self.buffer.size >= samples_per_chunk:
            chunk = self.buffer[:samples_per_chunk]
            self.buffer = self.buffer[samples_per_chunk:]

            wav_bytes = np_to_wav_bytes(chunk, self.SAMPLE_RATE)
            save_segment_to_session(wav_bytes)

            try:
                if backend.startswith("OpenAI"):
                    text = transcribe_bytes_openai(wav_bytes)
                else:
                    # For live mode, still use OpenAI by default (local streaming is non-trivial).
                    # You can switch to local after recording finishes.
                    text = transcribe_bytes_openai(wav_bytes)
                if text:
                    st.session_state.rolling_transcript.append(text)
            except Exception as e:
                st.session_state.rolling_transcript.append(f"[Transcribe error: {e}]")

        return frame

# WebRTC TURN configuration (use a public STUN; add your own TURN for corporate networks)
RTC_CFG = RTCConfiguration(
    {"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]}
)

# -------------------- LAYOUT --------------------
tab_live, tab_record, tab_upload, tab_minutes = st.tabs(["🎙️ Live", "⏺️ Record/Save", "📤 Upload", "🧾 Minutes"])

with tab_live:
    st.subheader("Live Transcription (rolling)")
    st.info("Press **Start** below. You’ll see rolling text every ~5 seconds. \
For the most accurate minutes, also click **Save full recording** in the Record tab and transcribe the final file.")

    ctx = webrtc_streamer(
        key="live",
        mode=WebRtcMode.SENDONLY,
        audio_receiver_size=256,
        rtc_configuration=RTC_CFG,
        media_stream_constraints={"audio": True, "video": False},
        audio_processor_factory=AudioChunker,
    )

    live_placeholder = st.empty()
    with live_placeholder.container():
        st.write("**Rolling transcript:**")
        st.write("")

    # Continuously show rolling transcript while session active
    if ctx.state.playing:
        # Render loop (works because Streamlit reruns on state updates)
        st.write("\n".join(st.session_state.rolling_transcript[-50:]))

with tab_record:
    st.subheader("Record meeting audio (optional)")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save full recording as WAV"):
            wav_bytes = export_full_audio_wav()
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button(
                "Download recording",
                data=wav_bytes,
                file_name=f"meeting_{ts}.wav",
                mime="audio/wav",
            )
    with col2:
        lang = st.text_input("Transcription language hint (optional, e.g., 'en' or 'zh')", value="")
        if st.button("📝 Transcribe saved recording (best quality)"):
            # Export and transcribe the combined audio from live session
            wav_bytes = export_full_audio_wav()
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tf:
                tf.write(wav_bytes)
                temp_path = tf.name
            with st.spinner("Transcribing final recording..."):
                if backend.startswith("OpenAI"):
                    text = transcribe_file_openai(temp_path, language=lang or None)
                else:
                    text = transcribe_file_faster(temp_path, language=lang or None)
            st.session_state.full_transcript = text
            st.success("Transcription complete (from recording). See the Minutes tab to summarize.")

with tab_upload:
    st.subheader("Upload audio for transcription")
    up = st.file_uploader("Upload audio (wav/mp3/m4a)", type=["wav", "mp3", "m4a"])
    lang2 = st.text_input("Language hint (optional)", key="u_lang")
    if up and st.button("Transcribe uploaded file"):
        with tempfile.NamedTemporaryFile(suffix="."+up.name.split(".")[-1], delete=False) as tf:
            tf.write(up.read())
            up_path = tf.name
        with st.spinner("Transcribing uploaded audio..."):
            if backend.startswith("OpenAI"):
                text = transcribe_file_openai(up_path, language=lang2 or None)
            else:
                text = transcribe_file_faster(up_path, language=lang2 or None)
        st.session_state.full_transcript = text
        st.success("Transcription complete. See the Minutes tab to summarize.")

with tab_minutes:
    st.subheader("Transcript & Minutes")

    # Compose best-available transcript
    rolling_text = " ".join(st.session_state.rolling_transcript).strip()
    full_text = st.session_state.get("full_transcript", "").strip()
    transcript_source = st.radio(
        "Choose transcript source",
        ["Rolling (live chunks)", "Final recording/Upload (recommended)"],
        index=1 if full_text else 0
    )
    if transcript_source.startswith("Rolling"):
        transcript = rolling_text
    else:
        transcript = full_text or rolling_text

    st.text_area("Transcript (editable)", transcript, height=240, key="__transcript_editor")

    colA, colB, colC = st.columns([1,1,1])
    with colA:
        if st.button("🧾 Generate Minutes"):
            if not st.session_state["__transcript_editor"].strip():
                st.warning("Transcript is empty.")
            else:
                with st.spinner("Summarizing to minutes..."):
                    try:
                        minutes = generate_minutes_llm(st.session_state["__transcript_editor"], minutes_sections or [])
                        st.session_state.minutes_text = minutes
                        st.success("Minutes generated.")
                    except Exception as e:
                        st.error(f"Minutes generation failed: {e}")

    if st.session_state.minutes_text:
        st.markdown("### Minutes")
        st.markdown(st.session_state.minutes_text)

        txt = st.session_state.minutes_text.encode("utf-8")
        st.download_button("⬇️ Download .txt", data=txt, file_name="meeting_minutes.txt", mime="text/plain")

        if HAVE_DOCX:
            doc = Document()
            doc.add_heading("Meeting Minutes", 0)
            for line in st.session_state.minutes_text.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("⬇️ Download .docx", data=buf.getvalue(), file_name="meeting_minutes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Install python-docx to enable .docx export:  pip install python-docx")
