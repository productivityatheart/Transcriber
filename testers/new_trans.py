# app.py — faster-whisper (STT) + Gemini (minutes)
import os
import io
import time
import tempfile
from datetime import datetime

import streamlit as st
import streamlit_webrtc
from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration, AudioProcessorBase
import av
import numpy as np
from pydub import AudioSegment

# ---- Transcription (local) ----
try:
    from faster_whisper import WhisperModel
    HAVE_FASTER = True
except Exception:
    HAVE_FASTER = False

# ---- Minutes export ----
try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# ---- Gemini for minutes ----
import google.generativeai as genai

@st.cache_resource
def configure_gemini() -> bool:
    # Try Streamlit secrets first, then environment variable
    key = (getattr(st, "secrets", {}).get("GEMINI_API_KEY")
           if hasattr(st, "secrets") else None)
    key = key or os.getenv("GEMINI_API_KEY")
    if not key:
        return False
    genai.configure(api_key=key)
    return True

gemini_ready = configure_gemini()

with st.sidebar:
    if gemini_ready:
        st.caption("✅ Gemini is configured from server secrets.")
    else:
        st.caption("⚠️ No server-side Gemini key found. Set .streamlit/secrets.toml or GEMINI_API_KEY.")

# ===== Cached Whisper model (shared across sessions) =====
@st.cache_resource(show_spinner="Loading faster-whisper model...")
def get_whisper(model_size: str, device: str, compute_type: str) -> WhisperModel:
    # device must be a string ("auto", "cpu", "cuda")
    dev = device if device != "auto" else "auto"
    return WhisperModel(model_size, device=dev, compute_type=compute_type)

# -------------------- UI SETUP --------------------
st.set_page_config(page_title="AI Meeting Minutes Generator", page_icon="📝", layout="wide")
st.title("📝 AI Meeting Minutes Generator")

with st.sidebar:
    st.header("⚙️ Settings")

    # We keep only local faster-whisper since no OpenAI key
    backend = st.selectbox("Transcription backend", ["Local faster-whisper (offline)"], index=0)

    # faster-whisper config
    model_size = st.selectbox("faster-whisper model", ["base", "small", "medium", "large-v3"], index=0)
    compute_type = st.selectbox("Compute type", ["int8", "int8_float16", "float16", "float32"], index=0)
    device = st.selectbox("Device", ["auto", "cpu", "cuda"], index=0)

    if HAVE_FASTER and st.button("Load / Warm up model"):
        _ = get_whisper(model_size, device, compute_type)   # cache + share
        st.success("faster-whisper model loaded and cached for all sessions.")

    st.divider()
    st.subheader("Minutes format")
    minutes_sections = st.multiselect(
        "Include sections",
    ["Agenda", "Key Points", "Action Items", "Summary"],
        default=["Agenda", "Key Points", "Action Items", "Summary"]
    )
    st.caption("Tip: live mode sends short rolling chunks; for best quality also transcribe the final recording.")
    
# -------------------- SESSION STATE --------------------
if "rolling_transcript" not in st.session_state:
    st.session_state.rolling_transcript = []
if "full_audio" not in st.session_state:
    st.session_state.full_audio = AudioSegment.silent(duration=0)
if "segments" not in st.session_state:
    st.session_state.segments = []  # [(start, end, text)]
if "minutes_text" not in st.session_state:
    st.session_state.minutes_text = ""

# -------------------- TRANSCRIBERS (use cached model) --------------------
def transcribe_chunk_faster(audio_np: np.ndarray, sample_rate: int, language: str | None = None) -> str:
    model = get_whisper(model_size, device, compute_type)
    segments, _ = model.transcribe(audio_np, language=language, vad_filter=True, sample_rate=sample_rate)
    return " ".join(s.text for s in segments).strip()

def transcribe_file_faster(path: str, language: str | None = None) -> str:
    model = get_whisper(model_size, device, compute_type)
    segments, _ = model.transcribe(path, language=language, vad_filter=True)
    out = []
    st.session_state.segments = []
    for seg in segments:
        out.append(seg.text)
        st.session_state.segments.append((seg.start, seg.end, seg.text))
    return " ".join(out).strip()

# -------------------- UTILITIES --------------------
def float_to_pcm16(audio: np.ndarray) -> np.ndarray:
    audio = np.clip(audio, -1.0, 1.0)
    return (audio * 32767.0).astype(np.int16)

def np_to_wav_bytes(audio_np: np.ndarray, sample_rate: int) -> bytes:
    pcm16 = float_to_pcm16(audio_np)
    seg = AudioSegment(data=pcm16.tobytes(), sample_width=2, frame_rate=sample_rate, channels=1)
    buf = io.BytesIO()
    seg.export(buf, format="wav")
    return buf.getvalue()

def save_segment_to_session(seg_bytes: bytes):
    new = AudioSegment.from_file(io.BytesIO(seg_bytes), format="wav")
    st.session_state.full_audio += new

def export_full_audio_wav() -> bytes:
    buf = io.BytesIO()
    st.session_state.full_audio.export(buf, format="wav")
    return buf.getvalue()

# -------------------- MINUTES GENERATOR (Gemini) --------------------
MINUTES_PROMPT = """You are an expert meeting minutes taker.
Given the raw transcript below, produce concise, structured minutes.
Use bullet points. Be specific with owners and dates when mentioned.
Sections to include (only output ones requested, in this order):
{sections}

Transcript:
\"\"\" 
{transcript}
\"\"\""""

def generate_minutes_gemini(transcript: str, sections: list[str]) -> str:
    if not gemini_ready:
        raise RuntimeError("Gemini is not configured on the server.")
    prompt = MINUTES_PROMPT.format(
        sections="\n".join(f"- {s}" for s in sections),
        transcript=transcript
    )
    model = genai.GenerativeModel("gemini-1.5-flash")
    resp = model.generate_content(prompt)
    return (resp.text or "").strip()


# -------------------- LIVE MODE (WebRTC) --------------------
class AudioChunker(AudioProcessorBase):
    SAMPLE_RATE = 16000
    CHANNELS = 1
    CHUNK_SECONDS = 5
    def __init__(self) -> None:
        self.buffer = np.zeros(0, dtype=np.float32)
    def recv_audio(self, frame: av.AudioFrame) -> av.AudioFrame:
        buf = frame.to_ndarray().astype(np.float32) / 32768.0  # int16 -> float
        if buf.ndim > 1:
            buf = buf.mean(axis=0)
        self.buffer = np.concatenate([self.buffer, buf])
        samples_per_chunk = int(self.SAMPLE_RATE * self.CHUNK_SECONDS)
        while self.buffer.size >= samples_per_chunk:
            chunk = self.buffer[:samples_per_chunk]
            self.buffer = self.buffer[samples_per_chunk:]
            wav_bytes = np_to_wav_bytes(chunk, self.SAMPLE_RATE)
            save_segment_to_session(wav_bytes)
            try:
                text = transcribe_chunk_faster(chunk, self.SAMPLE_RATE)
                if text:
                    st.session_state.rolling_transcript.append(text)
            except Exception as e:
                st.session_state.rolling_transcript.append(f"[Transcribe error: {e}]")
        return frame

RTC_CFG = RTCConfiguration({"iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]})

# -------------------- LAYOUT --------------------
tab_live, tab_record, tab_upload, tab_minutes = st.tabs(["🎙️ Live", "⏺️ Record/Save", "📤 Upload", "🧾 Minutes"])

with tab_live:
    st.subheader("Live Transcription (rolling)")
    st.info("Press **Start**. You’ll see rolling text every ~5 seconds. For best results, also transcribe the final recording.")
    ctx = webrtc_streamer(
        key=f"live-{st.session_state.get('webrtc_key', 0)}",
        mode=WebRtcMode.SENDONLY,
        audio_receiver_size=256,
        rtc_configuration=RTC_CFG,
        media_stream_constraints={"audio": True, "video": False},
        audio_processor_factory=AudioChunker,
    )
    if st.button("Retry mic permission"):
        st.session_state['webrtc_key'] = st.session_state.get('webrtc_key', 0) + 1
        st.experimental_rerun()
    live_placeholder = st.empty()
    with live_placeholder.container():
        st.write("**Rolling transcript:**")
        st.write("")
    if ctx.state.playing:
        st.write("\n".join(st.session_state.rolling_transcript[-50:]))

with tab_record:
    st.subheader("Record meeting audio (optional)")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save full recording as WAV"):
            wav_bytes = export_full_audio_wav()
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button("Download recording", data=wav_bytes, file_name=f"meeting_{ts}.wav", mime="audio/wav")
    with col2:
        lang = st.text_input("Transcription language hint (optional, e.g., 'en' or 'zh')", value="")
        if st.button("📝 Transcribe saved recording (best quality)"):
            wav_bytes = export_full_audio_wav()
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tf:
                tf.write(wav_bytes)
                temp_path = tf.name
            with st.spinner("Transcribing final recording..."):
                text = transcribe_file_faster(temp_path, language=lang or None)
            st.session_state.full_transcript = text
            st.success("Transcription complete (from recording). See the Minutes tab to summarize.")

with tab_upload:
    st.subheader("Upload audio for transcription")
    up = st.file_uploader("Upload audio (wav/mp3/m4a)", type=["wav", "mp3", "m4a"])
    lang2 = st.text_input("Language hint (optional)", key="u_lang")
    if up and st.button("Transcribe uploaded file"):
        with tempfile.NamedTemporaryFile(suffix="."+up.name.split(".")[-1], delete=False) as tf:
            tf.write(up.read())
            up_path = tf.name
        with st.spinner("Transcribing uploaded audio..."):
            text = transcribe_file_faster(up_path, language=lang2 or None)
        st.session_state.full_transcript = text
        st.success("Transcription complete. See the Minutes tab to summarize.")

with tab_minutes:
    st.subheader("Transcript & Minutes")
    rolling_text = " ".join(st.session_state.rolling_transcript).strip()
    full_text = st.session_state.get("full_transcript", "").strip()

    transcript_source = st.radio(
        "Choose transcript source",
        ["Rolling (live chunks)", "Final recording/Upload (recommended)"],
        index=1 if full_text else 0
    )
    transcript = rolling_text if transcript_source.startswith("Rolling") else (full_text or rolling_text)

    st.text_area("Transcript (editable)", transcript, height=240, key="__transcript_editor")

    colA, _, _ = st.columns([1,1,1])
    with colA:
        if st.button("🧾 Generate Minutes"):
            if not st.session_state["__transcript_editor"].strip():
                st.warning("Transcript is empty.")
            elif not gemini_ready:
                st.error("Gemini is not configured. Ask the admin to set `.streamlit/secrets.toml` or `GEMINI_API_KEY`.")
            else:
                with st.spinner("Summarizing to minutes (Gemini)..."):
                    try:
                        minutes = generate_minutes_gemini(
                            st.session_state["__transcript_editor"], 
                            minutes_sections or []
                        )
                        st.session_state.minutes_text = minutes
                        st.success("Minutes generated.")
                    except Exception as e:
                        st.error(f"Minutes generation failed: {e}")
    if st.session_state.minutes_text:
        st.markdown("### Minutes")
        st.markdown(st.session_state.minutes_text)
        txt = st.session_state.minutes_text.encode("utf-8")
        st.download_button("⬇️ Download .txt", data=txt, file_name="meeting_minutes.txt", mime="text/plain")
        if HAVE_DOCX:
            doc = Document()
            doc.add_heading("Meeting Minutes", 0)
            for line in st.session_state.minutes_text.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("⬇️ Download .docx", data=buf.getvalue(),
                               file_name="meeting_minutes.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.caption("Install python-docx to enable .docx export:  pip install python-docx")
