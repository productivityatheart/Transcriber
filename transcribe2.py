# transcribe.py — Record (browser) -> Transcribe (faster-whisper) -> Minutes (Gemini) -> Download ZIP

import os
import io
import re
import tempfile
import zipfile
from datetime import datetime
from typing import Any
import time, hashlib
import streamlit as st
import numpy as np
from pydub import AudioSegment
from audio_recorder_streamlit import audio_recorder
from docx import Document
import google.generativeai as genai

# ---------- Models / Keys ----------
try:
    from faster_whisper import WhisperModel
    HAVE_FASTER = True
except Exception as _FW_IMPORT_ERR:
    WhisperModel = None
    HAVE_FASTER = False

@st.cache_resource
def configure_gemini() -> bool:
    key = (getattr(st, "secrets", {}).get("GEMINI_API_KEY") if hasattr(st, "secrets") else None)
    key = key or os.getenv("GEMINI_API_KEY")
    if not key:
        return False
    genai.configure(api_key=key)
    return True

gemini_ready = configure_gemini()

@st.cache_resource(show_spinner="Loading faster-whisper model...")
def get_whisper(model_size: str = "base", device: str = "auto", compute_type: str = "int8") -> Any:
    if WhisperModel is None:
        raise RuntimeError(
            "faster-whisper is not available on the server. "
            f"Original import error: {_FW_IMPORT_ERR!r}"
        )
    dev = device if device in {"cpu", "cuda"} else "auto"
    return WhisperModel(model_size, device=dev, compute_type=compute_type)

# ---------- UI ----------
st.set_page_config(page_title="AI Meeting Minutes Generator", page_icon="📝", layout="wide")
st.title("📝 AI Meeting Minutes Generator")

with st.sidebar:
    st.caption("🧠 STT: faster-whisper (local) • Minutes: Gemini")
    st.caption("✅ Gemini configured from server secrets."
               if gemini_ready else "⚠️ Gemini key missing (.streamlit/secrets.toml or GEMINI_API_KEY).")
    title = st.text_input("Title (used for filenames)", placeholder="Weekly Sales Sync", value="")

# ---------- Session ----------
if "full_audio" not in st.session_state:
    st.session_state.full_audio = AudioSegment.silent(duration=0)
if "final_transcript" not in st.session_state:
    st.session_state.final_transcript = ""
if "minutes" not in st.session_state:
    st.session_state.minutes = ""
if "_user_prompt" not in st.session_state:
    st.session_state._user_prompt = "Create clean minutes: key points, action details, and a short summary."

# ---------- Helpers ----------
def export_full_audio_wav() -> bytes:
    buf = io.BytesIO()
    st.session_state.full_audio.export(buf, format="wav")
    return buf.getvalue()

def transcribe_audiosegment_with_fw(seg: AudioSegment, language: str | None = None) -> str:
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tf:
        seg.export(tf.name, format="wav")
        model = get_whisper("base", "auto", "int8")
        segments, _ = model.transcribe(tf.name, language=language, vad_filter=True)
        return " ".join(s.text for s in segments).strip()

def generate_minutes_with_gemini(transcript: str, user_prompt: str) -> str:
    if not gemini_ready:
        raise RuntimeError("Gemini is not configured on the server.")
    prompt = f"""
You are an expert meeting minutes writer.

User instruction:
{user_prompt}

Use the transcript below to produce concise, business-ready minutes.
- Bullet points where appropriate.
- Capture key points, decisions, action items (owners/dates if present), and a brief summary.
- If the user asked for a different structure, follow it.

Transcript:
\"\"\"{transcript}\"\"\""""
    model = genai.GenerativeModel("gemini-1.5-flash")
    resp = model.generate_content(prompt)
    return (getattr(resp, "text", "") or "").strip()

def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_-]+", "-", s)
    return s[:60] or "meeting"

# ---------- 1) Record ----------
st.subheader("1) Record")
st.caption("Click **Start**, speak, then click **Stop**. Takes append together.")

ss = st.session_state
# session state defaults
ss.setdefault("full_audio", AudioSegment.silent(duration=0))
ss.setdefault("rec_active", False)            # currently recording?
ss.setdefault("rec_key", "rec-stable")        # DO NOT change while recording
ss.setdefault("rec_started_at", None)         # datetime when Start was pressed
ss.setdefault("_pending_take", None)          # last blob to finalize after Stop
ss.setdefault("chunks_count", 0)
ss.setdefault("last_blob_ts", None)
ss.setdefault("last_blob_digest", None)       # to avoid double-append

# --- Start / Stop buttons ---
colA, colB = st.columns(2)
with colA:
    if st.button("🎙️ Start", disabled=ss.rec_active, key="btn_start"):
        ss.rec_active = True
        ss.rec_started_at = datetime.now()
        st.rerun()
with colB:
    if st.button("⏹ Stop", disabled=not ss.rec_active, key="btn_stop"):
        ss.rec_active = False
        st.rerun()

# --- When actively recording: render ONLY the recorder and freeze the app ---
if ss.rec_active:
    elapsed = (datetime.now() - ss.rec_started_at).seconds if ss.rec_started_at else 0
    st.info(f"⏱ Recording… {elapsed}s  •  🎧 Session total: {len(ss.full_audio)/1000:.1f}s")

    wav_live = audio_recorder(
        text="",                   # we show our own Start/Stop buttons
        sample_rate=16000,
        recording_color="#e74c3c",
        neutral_color="#2c3e50",
        icon_name="microphone",
        icon_size="2x",
        key=ss.rec_key,            # stable key prevents widget reset mid-take
    )

    # Remember latest blob; we'll append it only after Stop
    if wav_live:
        ss._pending_take = wav_live
        ss.last_blob_ts = time.time()

    st.stop()  # CRITICAL: prevents any rerun that would kill the recording

# --- Not recording: if we have a pending blob from the last run, finalize it ---
if ss._pending_take:
    wav_bytes = ss._pending_take
    ss._pending_take = None

    # de-dupe in case the same bytes reappear on a rerun
    digest = hashlib.sha1(wav_bytes).hexdigest()
    if digest != ss.last_blob_digest:
        ss.last_blob_digest = digest

        st.audio(wav_bytes, format="audio/wav")
        seg = AudioSegment.from_file(io.BytesIO(wav_bytes), format="wav")
        ss.full_audio += seg
        ss.chunks_count += 1

        taken = len(seg) / 1000.0
        total = len(ss.full_audio) / 1000.0
        # show how long the timer was running for this take
        show_elapsed = 0
        if ss.rec_started_at:
            show_elapsed = (datetime.now() - ss.rec_started_at).seconds
        ss.rec_started_at = None

        st.success(
            f"Captured {taken:.2f}s (timer ~{show_elapsed}s). "
            f"Total in session: {total:.1f}s (takes: {ss.chunks_count})"
        )

# --- Quick test / clear ---
c1, c2 = st.columns(2)
with c1:
    if st.button("🔊 Test: play combined audio"):
        st.audio(export_full_audio_wav(), format="audio/wav")
with c2:
    if st.button("🧹 Clear session audio"):
        ss.full_audio = AudioSegment.silent(duration=0)
        ss.chunks_count = 0
        ss.rec_started_at = None
        ss._pending_take = None
        ss.last_blob_digest = None
        ss.last_blob_ts = None
        st.success("Cleared.")

st.divider()

# ---------- 2) Transcribe & Minutes ----------
st.subheader("2) Transcribe & Minutes")
colX, colY = st.columns([1, 2])

with colX:
    if st.button("📝 Transcribe & Generate"):
        if len(st.session_state.full_audio) == 0:
            st.warning("No audio recorded yet.")
        else:
            with st.spinner("Transcribing (faster-whisper)…"):
                st.session_state.final_transcript = transcribe_audiosegment_with_fw(
                    st.session_state.full_audio
                )
            if not gemini_ready:
                st.error("Gemini is not configured (.streamlit/secrets.toml or GEMINI_API_KEY).")
            else:
                with st.spinner("Generating minutes (Gemini)…"):
                    st.session_state.minutes = generate_minutes_with_gemini(
                        st.session_state.final_transcript, st.session_state._user_prompt
                    )
            st.success("Done. See Preview and Download below.")

with colY:
    st.session_state._user_prompt = st.text_area(
        "Prompt for AI (what to do with the transcript)",
        value=st.session_state._user_prompt,
        height=120,
    )

# ---------- Preview ----------
st.subheader("Preview")
cA, cB = st.columns(2)
with cA:
    st.markdown("**Transcript**")
    st.text_area("transcript.txt", st.session_state.final_transcript, height=260)
with cB:
    st.markdown("**Minutes**")
    st.text_area("minutes.txt", st.session_state.minutes, height=260)

# ---------- 3) Download ZIP ----------
st.subheader("3) Download package (audio + transcript + minutes)")

def build_zip_package(title_str: str) -> bytes:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{slugify(title_str) or 'meeting'}_{ts}"
    mem = io.BytesIO()

    # DOCX minutes
    minutes_doc = Document()
    minutes_doc.add_heading("Meeting Minutes", level=0)
    for line in (st.session_state.minutes or "").splitlines():
        minutes_doc.add_paragraph(line if line.strip() else "")
    mbuf = io.BytesIO(); minutes_doc.save(mbuf); minutes_bytes = mbuf.getvalue()

    # DOCX transcript
    transcript_doc = Document()
    transcript_doc.add_heading("Transcript", level=0)
    for line in (st.session_state.final_transcript or "").splitlines():
        transcript_doc.add_paragraph(line if line.strip() else "")
    tbuf = io.BytesIO(); transcript_doc.save(tbuf); transcript_bytes = tbuf.getvalue()

    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{base}/audio.wav", export_full_audio_wav())
        zf.writestr(f"{base}/minutes.docx", minutes_bytes)
        zf.writestr(f"{base}/transcript.docx", transcript_bytes)
        meta = f"Title: {title_str or 'meeting'}\nCreated: {ts}\n"
        zf.writestr(f"{base}/meta.txt", meta.encode("utf-8"))

    mem.seek(0)
    return mem.getvalue()

pkg_disabled = (len(st.session_state.full_audio) == 0)
pkg_bytes = build_zip_package(title) if not pkg_disabled else None

st.download_button(
    "⬇️ Download ZIP",
    data=pkg_bytes if pkg_bytes else b"",
    file_name=f"{slugify(title) or 'meeting'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
    mime="application/zip",
    disabled=pkg_disabled,
)
