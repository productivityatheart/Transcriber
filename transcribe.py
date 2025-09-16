# transcribe.py — Continuous (WebRTC) Recorder → Transcribe (faster-whisper) → Minutes (Gemini) → Download ZIP
import os, io, re, time, wave, zipfile
from datetime import datetime, timedelta
from typing import Any, Optional
import contextlib
import shutil

import av
import numpy as np
import streamlit as st
from docx import Document
from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration, AudioProcessorBase
import google.generativeai as genai

# =========================
# Streamlit (must be first)
# =========================
st.set_page_config(page_title="AI Meeting Minutes Generator", page_icon="📝", layout="wide")

# =========================
# Settings
# =========================
OUTPUT_DIR = os.environ.get("REC_DIR", "/tmp/recordings")
os.makedirs(OUTPUT_DIR, exist_ok=True)

SAMPLE_RATE = 48000   # browser-native
CHANNELS = 1          # mono
SAMPLE_WIDTH = 2      # 16-bit PCM
ROTATE_EVERY_MIN = 10

RTC_CFG = RTCConfiguration({
    "iceServers": [
        {"urls": ["stun:stun.l.google.com:19302"]},
        # Add TURN here if needed
    ]
})

# ---------- Session state (must be before webrtc_streamer) ----------
ss = st.session_state
ss.setdefault("session_id", f"sess_{int(time.time())}")
ss.setdefault("record_base", os.path.join(OUTPUT_DIR, ss.session_id))
ss.setdefault("webrtc_started_at", None)
ss.setdefault("parts", [])
ss.setdefault("merged_path", "")
ss.setdefault("transcript", "")
ss.setdefault("minutes", "")
ss.setdefault("writer", None)

record_base = ss["record_base"]
st.caption(f"Files will be saved as: `{record_base}_partNN.wav`")

# =========================
# Utilities
# =========================
def float_to_pcm16(x: np.ndarray) -> np.ndarray:
    x = np.clip(x, -1.0, 1.0)
    return (x * 32767.0).astype(np.int16)

def merge_wavs(part_files: list[str], out_path: str):
    """Merge 16-bit PCM WAVs with identical params."""
    if not part_files:
        raise RuntimeError("No part files to merge.")
    with wave.open(part_files[0], "rb") as w0:
        params = w0.getparams()
        chunks = [w0.readframes(w0.getnframes())]
    for p in part_files[1:]:
        with wave.open(p, "rb") as wi:
            if (wi.getnchannels(), wi.getsampwidth(), wi.getframerate()) != (params.nchannels, params.sampwidth, params.framerate):
                raise ValueError(f"Incompatible WAV params in {p}")
            chunks.append(wi.readframes(wi.getnframes()))
    with wave.open(out_path, "wb") as wo:
        wo.setparams(params)
        for fr in chunks:
            wo.writeframes(fr)

def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_-]+", "-", s)
    return s[:60] or "meeting"

# =========================
# Audio writer (WebRTC worker)
# =========================
class RollingWavWriter(AudioProcessorBase):
    """
    Buffer incoming audio as PCM16 in memory and write fully-closed WAV files
    on rotation/close. Also track received frames for debugging.
    """
    def __init__(self, base_path: str, rotate_every_min: int = 10) -> None:
        self.base_path = base_path
        self.rotate_every = timedelta(minutes=rotate_every_min)
        self.part_idx = 0
        self.last_rotate: Optional[datetime] = None
        self.part_paths: list[str] = []
        self.frames_received: int = 0
        self.bytes_buffered: int = 0

        os.makedirs(os.path.dirname(base_path), exist_ok=True)
        self._buf = bytearray()
        self._cur_path: Optional[str] = None
        self._start_new_part()

    def _write_pcm16_wav(self, out_path: str, pcm_bytes: bytes, force: bool = False):
        # If force=True and pcm is empty, write 0.5s of silence so the WAV is valid
        if not pcm_bytes and force:
            pcm_bytes = b"\x00" * int(0.5 * SAMPLE_RATE * CHANNELS * SAMPLE_WIDTH)
        if not pcm_bytes:
            return
        with wave.open(out_path, "wb") as w:
            w.setnchannels(CHANNELS)
            w.setsampwidth(SAMPLE_WIDTH)
            w.setframerate(SAMPLE_RATE)
            w.writeframes(pcm_bytes)

    def _flush_current(self, force: bool = False):
        if self._cur_path is not None:
            self._write_pcm16_wav(self._cur_path, bytes(self._buf), force=force)

    def _start_new_part(self):
        # Flush previous (closed WAV on disk)
        self._flush_current(force=False)
        self.part_idx += 1
        self.last_rotate = datetime.utcnow()
        self._cur_path = f"{self.base_path}_part{self.part_idx:02d}.wav"
        if self._cur_path not in self.part_paths:
            self.part_paths.append(self._cur_path)
        self._buf = bytearray()

    def recv_audio(self, frame: av.AudioFrame) -> av.AudioFrame:
        arr = frame.to_ndarray()
        if arr.ndim == 2:
            arr = arr.mean(axis=0 if arr.shape[0] <= arr.shape[1] else 1)

        if np.issubdtype(arr.dtype, np.integer):
            maxv = np.iinfo(arr.dtype).max
            audio = arr.astype(np.float32) / float(maxv)
        else:
            audio = arr.astype(np.float32)

        pcm16 = float_to_pcm16(audio)
        b = pcm16.tobytes()
        self._buf += b
        self.frames_received += 1
        self.bytes_buffered += len(b)

        if self.last_rotate and (datetime.utcnow() - self.last_rotate) >= self.rotate_every:
            self._start_new_part()

        return frame

    def close(self):
        # Force flush ensures a valid WAV even if nothing was captured
        self._flush_current(force=True)
        self._buf = bytearray()

# =========================
# Models / Keys
# =========================
_FW_IMPORT_ERR = None  # ensure it's always defined

try:
    from faster_whisper import WhisperModel
    HAVE_FASTER = True
except Exception as e:
    WhisperModel = None
    HAVE_FASTER = False
    _FW_IMPORT_ERR = e

@st.cache_resource
def configure_gemini() -> bool:
    key = (getattr(st, "secrets", {}).get("GEMINI_API_KEY") if hasattr(st, "secrets") else None)
    key = key or os.getenv("GEMINI_API_KEY")
    if not key:
        return False
    genai.configure(api_key=key)
    return True

gemini_ready = configure_gemini()

@st.cache_resource(show_spinner="Loading faster-whisper model...")
def get_whisper(model_size: str = "base", device: str = "auto", compute_type: str = "int8") -> Any:
    if not HAVE_FASTER or WhisperModel is None:
        msg = "faster-whisper is not available. Install with `pip install faster-whisper`."
        if _FW_IMPORT_ERR is not None:
            msg += f" Import error: {repr(_FW_IMPORT_ERR)}"
        raise RuntimeError(msg)
    dev = device if device in {"cpu", "cuda"} else "auto"
    return WhisperModel(model_size, device=dev, compute_type=compute_type)

def transcribe_file_faster(path: str, language: Optional[str] = None) -> str:
    model = get_whisper("base", "auto", "int8")
    segments, _ = model.transcribe(path, language=language, vad_filter=True)
    return " ".join(s.text for s in segments).strip()

def generate_minutes_gemini(transcript: str, prompt: str) -> str:
    """
    Use Gemini to turn a transcript into concise meeting minutes.
    Assumes `gemini_ready` and `genai` are available.
    """
    if not transcript or not transcript.strip():
        return "Transcript is empty—nothing to summarize."
    if not gemini_ready:
        raise RuntimeError("Gemini is not configured (.streamlit/secrets.toml or GEMINI_API_KEY).")

    sys_prompt = f"""
You are an expert meeting minutes writer.

User instruction:
{prompt}

Use the transcript below to produce concise, business-ready minutes.
- Use clear bullets where appropriate.
- Capture key points, decisions, action items (owners/dates if present), and a brief summary.
- If the user asked for a different structure, follow it.

Transcript:
\"\"\"{transcript}\"\"\""""

    model = genai.GenerativeModel("gemini-1.5-flash")
    resp = model.generate_content(sys_prompt)
    return (getattr(resp, "text", "") or "").strip()

# =========================
# UI
# =========================
st.title("📝 AI Meeting Minutes Generator")

with st.sidebar:
    st.caption("🎙️ Recorder: WebRTC → rotating WAV on disk")
    st.caption("🧠 STT: faster-whisper • Minutes: Gemini")
    st.caption("✅ Gemini via server secrets" if gemini_ready else "⚠️ Add GEMINI_API_KEY")
    meeting_title = st.text_input("Title (for filenames)", placeholder="Weekly Sales Sync", value="")
    user_prompt = st.text_area(
        "Prompt for minutes",
        value="Create clean minutes: key points, decisions, action items (with owners/dates), and a brief summary.",
        height=120,
    )

# ---------- 1) Start / Stop ----------
BASE_RECORD_PATH = ss["record_base"]

# Keep a handle so we can close cleanly when the stream stops
_writer_handle = {"w": None}

def _writer_factory(base=BASE_RECORD_PATH):
    # DO NOT touch st.session_state here (worker thread)
    w = RollingWavWriter(base, rotate_every_min=ROTATE_EVERY_MIN)
    _writer_handle["w"] = w
    return w

ctx = webrtc_streamer(
    key="long-recorder",
    mode=WebRtcMode.SENDONLY,
    rtc_configuration=RTC_CFG,
    media_stream_constraints={"audio": True, "video": False},
    audio_processor_factory=_writer_factory,
)

# Start/stop timer + ensure file handles get closed
if ctx and ctx.state.playing:
    if ss.get("webrtc_started_at") is None:
        ss["webrtc_started_at"] = datetime.utcnow()
    elapsed = (datetime.utcnow() - ss["webrtc_started_at"]).seconds
    st.info(f"⏱ Recording… {elapsed//60:02d}:{elapsed%60:02d}")
else:
    w = _writer_handle.get("w")
    if w is not None:
        try:
            w.close()  # forces a valid WAV even if no frames arrived
        finally:
            _writer_handle["w"] = None
    if ss.get("webrtc_started_at") is not None:
        st.success("Stopped recording.")
    ss["webrtc_started_at"] = None

# ===== Helpers (defined before they are used) =====
def is_valid_wav(path: str) -> bool:
    try:
        with contextlib.closing(wave.open(path, "rb")) as r:
            _ = (r.getnchannels(), r.getsampwidth(), r.getframerate(), r.getnframes())
        return True
    except Exception:
        return False

def list_parts(base: str) -> list[str]:
    d = os.path.dirname(base)
    prefix = os.path.basename(base) + "_part"
    if not os.path.isdir(d):
        return []
    files = sorted(
        os.path.join(d, f) for f in os.listdir(d)
        if f.startswith(prefix) and f.endswith(".wav")
    )
    # keep only readable/closed parts
    return [p for p in files if is_valid_wav(p)]

# --- Debug panel: show write status & auto-scan parts when stopped ---
st.divider()
st.subheader("Recorder status")

writer = _writer_handle.get("w")
if ctx and ctx.state.playing:
    if writer:
        st.caption(f"frames_received={writer.frames_received:,} • bytes_buffered={writer.bytes_buffered:,}")
        if writer.part_paths:
            cur_part = writer.part_paths[-1]
            size = os.path.getsize(cur_part) if os.path.exists(cur_part) else 0
            st.code("\n".join(os.path.basename(p) for p in writer.part_paths), language="text")
            st.caption(f"📄 Current part: `{os.path.basename(cur_part)}` • Size: {size:,} bytes")
    else:
        st.caption("⏳ Waiting for first audio frames… (check mic permission)")
    total, used, free = shutil.disk_usage(os.path.dirname(ss["record_base"]))
    st.caption(f"💾 Free space: {free // (1024*1024)} MB")
else:
    rec_dir = os.path.dirname(ss["record_base"])
    if os.path.isdir(rec_dir):
        st.code("\n".join(sorted(os.listdir(rec_dir))), language="text")
    # auto-scan closed parts (valid WAVs only)
    ss["parts"] = list_parts(ss["record_base"])
    st.caption(f"✅ Found {len(ss['parts'])} closed parts in `{rec_dir}`")
    if st.button("🔧 Force finalize current part"):
        if writer:
            writer.close()
        ss["parts"] = list_parts(ss["record_base"])
        st.success(f"Finalized. Now {len(ss['parts'])} closed parts.")

# ========== After you click Stop ==========
is_recording = bool(ctx and ctx.state.playing)
if is_recording:
    st.warning("Recording in progress — stop first before scanning/previewing/merging.")

st.divider()
st.subheader("After you click Stop")

colA, colB = st.columns([1, 1])

with colA:
    if st.button("🔎 Scan parts", disabled=is_recording):
        ss["parts"] = list_parts(ss["record_base"])
        st.json(ss.get("parts", []))

with colB:
    if st.button("🧩 Merge parts to single WAV", disabled=is_recording or not ss.get("parts")):
        parts_list = list_parts(ss["record_base"])  # re-scan closed, valid parts
        if not parts_list:
            st.warning("No closed audio parts to merge. Click Stop, then Scan parts.")
        else:
            merged = f"{ss['record_base']}_merged.wav"
            merge_wavs(parts_list, merged)
            ss["merged_path"] = merged
            st.success(f"Merged → {merged}")

    if ss.get("merged_path"):
        with open(ss["merged_path"], "rb") as f:
            st.download_button(
                "⬇️ Download merged WAV",
                data=f.read(),
                file_name=os.path.basename(ss["merged_path"]),
                mime="audio/wav",
            )

st.markdown("**Or upload an audio file instead (fallback):**")
upl = st.file_uploader("Accepts WAV/MP3/M4A/etc.", type=["wav", "mp3", "m4a", "aac", "flac", "ogg"], accept_multiple_files=False)
if upl is not None:
    up_path = os.path.join(OUTPUT_DIR, f"upload_{int(time.time())}_{slugify(upl.name)}")
    with open(up_path, "wb") as f:
        f.write(upl.read())
    ss["merged_path"] = up_path
    st.success(f"Uploaded and ready: {os.path.basename(up_path)}")

# ========== Transcribe & Minutes ==========
st.divider()
st.subheader("2) Transcribe & Minutes")

col1, col2 = st.columns([1, 2])

with col1:
    if st.button("📝 Transcribe & Generate", disabled=is_recording):
        if not ss.get("merged_path"):
            parts_list = list_parts(ss["record_base"])
            if not parts_list:
                st.warning("No audio available. Use the upload option above, or record then Stop.")
                st.stop()
            merged = f"{ss['record_base']}_merged.wav"
            merge_wavs(parts_list, merged)
            ss["merged_path"] = merged

        try:
            with st.spinner("Transcribing with faster-whisper…"):
                ss["transcript"] = transcribe_file_faster(ss["merged_path"])
        except Exception as e:
            st.error(f"Transcription failed: {e}")
            st.stop()

        if not gemini_ready:
            st.error("Gemini is not configured (.streamlit/secrets.toml or GEMINI_API_KEY).")
        else:
            with st.spinner("Generating minutes (Gemini)…"):
                ss["minutes"] = generate_minutes_gemini(ss["transcript"], user_prompt)

        st.success("Done. See Preview and Download below.")

with col2:
    st.text_area("Transcript (preview)", ss.get("transcript", ""), height=180)
    st.text_area("Minutes (preview)", ss.get("minutes", ""), height=180)

# ========== Download ==========
st.divider()
st.subheader("3) Download package (audio + transcript + minutes)")

def build_zip(meeting_title: str) -> bytes:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{slugify(meeting_title) or 'meeting'}_{ts}"
    mem = io.BytesIO()

    wav_bytes = b""
    if ss.get("merged_path") and os.path.exists(ss["merged_path"]):
        with open(ss["merged_path"], "rb") as f:
            wav_bytes = f.read()

    # DOCX minutes
    minutes_doc = Document()
    minutes_doc.add_heading("Meeting Minutes", level=0)
    for line in (ss.get("minutes") or "").splitlines():
        minutes_doc.add_paragraph(line if line.strip() else "")
    mbuf = io.BytesIO()
    minutes_doc.save(mbuf)
    minutes_bytes = mbuf.getvalue()

    # DOCX transcript
    transcript_doc = Document()
    transcript_doc.add_heading("Transcript", level=0)
    for line in (ss.get("transcript") or "").splitlines():
        transcript_doc.add_paragraph(line if line.strip() else "")
    tbuf = io.BytesIO()
    transcript_doc.save(tbuf)
    transcript_bytes = tbuf.getvalue()

    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        if wav_bytes:
            zf.writestr(f"{base}/audio.wav", wav_bytes)
        zf.writestr(f"{base}/transcript.txt", (ss.get("transcript") or "").encode("utf-8"))
        zf.writestr(f"{base}/minutes.docx", minutes_bytes)
        zf.writestr(f"{base}/transcript.docx", transcript_bytes)
        meta = f"Title: {meeting_title or 'meeting'}\nCreated: {ts}\n"
        zf.writestr(f"{base}/meta.txt", meta.encode("utf-8"))

    mem.seek(0)
    return mem.getvalue()

zip_disabled = not ss.get("merged_path") or not os.path.exists(ss["merged_path"])
st.download_button(
    "⬇️ Download ZIP",
    data=(build_zip(meeting_title) if not zip_disabled else b""),
    file_name=f"{slugify(meeting_title) or 'meeting'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
    mime="application/zip",
    disabled=zip_disabled,
)
